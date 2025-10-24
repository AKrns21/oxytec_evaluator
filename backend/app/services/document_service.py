"""Document processing service for extracting text from various formats."""

import os
import base64
import hashlib
from pathlib import Path
from typing import Optional
import fitz  # PyMuPDF
import docx
import pandas as pd
from sqlalchemy import select
from app.db.session import AsyncSessionLocal
from app.models.database import Document
from app.utils.logger import get_logger
from app.utils.error_handler import handle_service_errors

logger = get_logger(__name__)


class DocumentService:
    """Service for extracting text from documents."""

    @handle_service_errors("document_extraction")
    async def extract_text(self, file_path: str, mime_type: Optional[str] = None) -> str:
        """
        Extract text from a document with caching support.

        Checks database cache first. If cached extraction exists and file hasn't changed,
        returns cached content. Otherwise, performs extraction and caches result.

        Supports: PDF, DOCX, TXT, CSV, XLSX, PNG, JPG/JPEG (via vision)

        Args:
            file_path: Path to the file
            mime_type: MIME type of the file (optional)

        Returns:
            Extracted text content
        """

        path = Path(file_path)

        # Check cache first
        cached_text = await self._get_cached_extraction(str(path))
        if cached_text:
            logger.info("extraction_cache_hit", file_path=str(path))
            return cached_text

        # Cache miss - perform extraction
        logger.info("extraction_cache_miss", file_path=str(path))

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower()

        # Extract text based on file type
        if extension == ".pdf":
            extracted_text = await self._extract_pdf(file_path)
        elif extension in [".docx", ".doc"]:
            extracted_text = await self._extract_docx(file_path)
        elif extension == ".txt":
            extracted_text = await self._extract_txt(file_path)
        elif extension in [".csv", ".xlsx", ".xls"]:
            extracted_text = await self._extract_spreadsheet(file_path)
        elif extension in [".png", ".jpg", ".jpeg"]:
            extracted_text = await self._extract_image(file_path)
        else:
            logger.warning("unsupported_file_type", extension=extension)
            extracted_text = f"[Unsupported file type: {extension}]"

        # Cache the extracted text
        await self._cache_extraction(str(path), extracted_text)

        return extracted_text

    @handle_service_errors("pdf_extraction")
    async def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF, fallback to vision for image-based PDFs."""
        doc = fitz.open(file_path)
        text_parts = []
        image_based_pages = []

        # First, try regular text extraction
        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            if text.strip():
                text_parts.append(f"--- Page {page_num} ---\n{text}")
            else:
                # No text found, this might be an image-based page
                image_based_pages.append((page_num, page))

        # If we got text from regular extraction, use it
        if text_parts:
            doc.close()
            return "\n\n".join(text_parts)

        # If no text was extracted, use vision to extract from images
        logger.info("pdf_is_image_based", file_path=file_path, pages=len(doc))

        # Extract using vision for image-based PDFs (in parallel)
        import asyncio
        from app.config import settings

        # Convert all pages to images first
        page_images = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
            img_bytes = pix.tobytes("png")
            page_images.append((page_num + 1, img_bytes))

        # Check if dual-model mode is enabled
        if settings.vision_use_dual_model:
            logger.info("vision_dual_model_enabled", pages=len(page_images))

            # Run both OpenAI and Gemini in parallel for all pages
            openai_tasks = [
                self._extract_text_from_image_with_openai(
                    img_bytes,
                    f"page_{page_num}"
                )
                for page_num, img_bytes in page_images
            ]

            gemini_tasks = [
                self._extract_text_from_image_with_vision(
                    img_bytes,
                    f"page_{page_num}"
                )
                for page_num, img_bytes in page_images
            ]

            # Execute both model batches in parallel
            openai_results, gemini_results = await asyncio.gather(
                asyncio.gather(*openai_tasks, return_exceptions=True),
                asyncio.gather(*gemini_tasks, return_exceptions=True)
            )

            # Blend results page by page
            vision_results = []
            for idx, (page_num, _) in enumerate(page_images):
                openai_result = openai_results[idx] if not isinstance(openai_results[idx], Exception) else f"[OpenAI error: {str(openai_results[idx])}]"
                gemini_result = gemini_results[idx] if not isinstance(gemini_results[idx], Exception) else f"[Gemini error: {str(gemini_results[idx])}]"

                # Blend and choose best result
                blended_result, chosen_model = self._blend_vision_results(
                    openai_result,
                    gemini_result,
                    page_num
                )
                vision_results.append(blended_result)
        else:
            # Single model mode (Gemini only)
            logger.info("vision_single_model_gemini", pages=len(page_images))

            # Extract text from all pages in parallel using Gemini only
            vision_tasks = [
                self._extract_text_from_image_with_vision(
                    img_bytes,
                    f"page_{page_num}"
                )
                for page_num, img_bytes in page_images
            ]

            vision_results = await asyncio.gather(*vision_tasks, return_exceptions=True)

        # Collect results - IMPORTANT: Include ALL pages to preserve page numbering
        vision_text_parts = []
        for idx, (page_num, _) in enumerate(page_images):
            result = vision_results[idx]
            if isinstance(result, Exception):
                logger.error("vision_page_failed", page=page_num, error=str(result))
                vision_text_parts.append(f"--- Page {page_num} ---\n[Vision extraction failed: {str(result)}]")
            elif result.strip():
                vision_text_parts.append(f"--- Page {page_num} ---\n{result}")
            else:
                # Include empty pages to preserve page numbering
                logger.warning("vision_page_empty", page=page_num)
                vision_text_parts.append(f"--- Page {page_num} ---\n[Vision API returned empty response - page may be blank or extraction failed]")

        doc.close()

        if vision_text_parts:
            logger.info("pdf_extracted_with_vision", pages=len(vision_text_parts))
            return "\n\n".join(vision_text_parts)
        else:
            logger.warning("pdf_no_content_extracted", file_path=file_path)
            return "[No text content could be extracted from this PDF]"

    @handle_service_errors("docx_extraction")
    async def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        doc = docx.Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Also extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)

        all_text = paragraphs
        if tables_text:
            all_text.append("\n--- Tables ---\n")
            all_text.extend(tables_text)

        return "\n".join(all_text)

    @handle_service_errors("txt_extraction")
    async def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding - this is a special case we want to handle
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()

    @handle_service_errors("image_extraction")
    async def _extract_image(self, file_path: str) -> str:
        """
        Extract text from image file (PNG, JPG, JPEG) using Claude's vision API.

        Args:
            file_path: Path to the image file

        Returns:
            Extracted text content from the image
        """
        logger.info("image_extraction_started", file_path=file_path)

        # Read image file as bytes
        with open(file_path, "rb") as f:
            image_bytes = f.read()

        # Determine image format from extension
        extension = Path(file_path).suffix.lower()
        if extension == ".png":
            media_type = "image/png"
        elif extension in [".jpg", ".jpeg"]:
            media_type = "image/jpeg"
        else:
            logger.warning("unsupported_image_format", extension=extension)
            return f"[Unsupported image format: {extension}]"

        # Extract text using vision API
        from anthropic import AsyncAnthropic
        from app.config import settings

        # Encode image to base64
        image_base64 = base64.b64encode(image_bytes).decode('utf-8')

        # Use Claude with vision
        client = AsyncAnthropic(api_key=settings.anthropic_api_key)

        response = await client.messages.create(
            model=settings.anthropic_model,
            max_tokens=4096,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": image_base64
                        }
                    },
                    {
                        "type": "text",
                        "text": """You are a document digitization specialist. Your task is to convert this document image into a comprehensive, structured JSON format that captures ALL content.

CRITICAL: Extract EVERYTHING visible - miss nothing. This is the only chance to capture this data.

Return a JSON object with this structure:

{
  "document_type": "string (email, technical_drawing, table, form, safety_data_sheet, measurement_report, process_flow_diagram, questionnaire, mixed)",
  "metadata": {
    "has_tables": boolean,
    "has_diagrams": boolean,
    "has_handwriting": boolean,
    "language": "string (de, en, mixed)",
    "page_number": "string or null if visible"
  },
  "content": {
    "body_text": "string (ALL text in document order. CRITICAL: Write headers WITH their paragraphs together, not separately! Format: HEADER\\n\\nparagraph content\\n\\nNEXT HEADER\\n\\nparagraph. Preserve document structure!)",
    "tables": [
      {
        "title": "string or null",
        "headers": ["column1", "column2", ...],
        "rows": [
          ["cell1", "cell2", ...],
          ["cell1", "cell2", ...]
        ]
      }
    ],
    "lists": [
      {
        "type": "bulleted or numbered",
        "items": ["item1", "item2", ...]
      }
    ],
    "key_value_pairs": [
      {"key": "string", "value": "string"}
    ],
    "diagrams_and_images": [
      {
        "type": "flow_diagram, chart, logo, signature, stamp, photo, technical_drawing",
        "description": "detailed description of what is shown",
        "labels_and_text": ["any text visible in/on the diagram"]
      }
    ],
    "signatures_and_stamps": [
      {
        "type": "signature or stamp",
        "text": "any readable text",
        "location": "top_right, bottom_left, etc."
      }
    ]
  },
  "quality_notes": "string (mention any unclear text, cut-off content, poor quality areas)"
}

EXTRACTION RULES:

1. **Tables**:
   - Extract EVERY row and column
   - Preserve exact values, units, symbols (%, ≤, ≥, -, ~)
   - If cells span multiple rows/columns, note this
   - Include table headers AND all data rows

2. **Text**:
   - Extract ALL paragraphs verbatim
   - Preserve line breaks between sections
   - Include page numbers, headers, footers
   - Capture email signatures, contact info

3. **Key-Value Pairs**:
   - Extract form fields: "Field Name: Value"
   - Extract measurement data: "Temperature: 45°C"
   - Extract parameters: "Flow Rate: 5000 m³/h"

4. **Preserve Formatting**:
   - Keep units exactly: m³/h, °C, mg/Nm³, %
   - Keep special characters: ≤, ≥, ±, ~, -, /, ×
   - Keep German umlauts: ä, ö, ü, ß

5. **Diagrams/Images**:
   - Describe what is shown (process flow, equipment layout, etc.)
   - Extract ALL labels, annotations, arrows, text from diagrams
   - Note connections between elements

6. **Don't Miss**:
   - Small print, footnotes, references
   - Handwritten notes or annotations
   - Stamps, signatures, logos with text
   - Section numbers, page numbers
   - CAS numbers, chemical formulas
   - Email headers, sender/recipient info

Return ONLY valid JSON. No markdown, no commentary."""
                    }
                ]
            }]
        )

        # Extract text from response
        if not response.content or len(response.content) == 0:
            logger.warning("vision_response_empty", file_path=file_path)
            return "[No content extracted from image]"

        content_block = response.content[0]
        if not hasattr(content_block, 'text'):
            logger.error("vision_response_no_text", file_path=file_path)
            return "[Vision response has no text content]"

        extracted_text = content_block.text

        logger.info(
            "image_extraction_success",
            file_path=file_path,
            length=len(extracted_text)
        )

        return extracted_text

    @handle_service_errors("spreadsheet_extraction")
    async def _extract_spreadsheet(self, file_path: str) -> str:
        """Extract and intelligently summarize spreadsheet data."""
        from app.config import settings

        extension = Path(file_path).suffix.lower()
        filename = Path(file_path).name

        if extension == ".csv":
            df = pd.read_csv(file_path)
            return await self._summarize_dataframe(df, filename)
        else:
            # Read all sheets from Excel file
            sheets_dict = pd.read_excel(file_path, sheet_name=None)

            # Handle single vs multi-sheet Excel
            if len(sheets_dict) == 1:
                # Single sheet - extract it directly
                sheet_name, df = list(sheets_dict.items())[0]
                return await self._summarize_dataframe(df, filename, sheet_name)
            else:
                # Multiple sheets - summarize each
                return await self._extract_multi_sheet_excel(sheets_dict, filename)

    async def _summarize_dataframe(
        self,
        df: pd.DataFrame,
        filename: str,
        sheet_name: Optional[str] = None
    ) -> str:
        """
        Intelligently summarize DataFrame based on content type.

        For measurement/numerical data (>50% numeric columns):
        - Provides statistical summary (min, max, mean, median, std)
        - Shows representative samples (first and last few rows)

        For text/categorical data:
        - Shows first N rows as preview

        This dramatically reduces token count for LLM processing while
        preserving critical information for feasibility analysis.
        """
        from app.config import settings

        # Build header
        summary_parts = []

        if sheet_name:
            summary_parts.append(f"Sheet: {sheet_name}")
        else:
            summary_parts.append(f"File: {filename}")

        summary_parts.extend([
            f"Dimensions: {len(df)} rows × {len(df.columns)} columns",
            f"Columns: {', '.join(str(col) for col in df.columns.tolist())}",
            ""
        ])

        # Detect if this is measurement/numerical data
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        numeric_ratio = len(numeric_cols) / len(df.columns) if len(df.columns) > 0 else 0

        # For very small datasets, just show the full data
        if len(df) <= 10:
            summary_parts.append("--- Complete Data ---")
            summary_parts.append(df.to_string(index=False))
            return "\n".join(summary_parts)

        # Strategy 1: Statistical summary for measurement data
        if numeric_ratio > settings.excel_statistical_threshold:
            logger.info(
                "using_statistical_summary",
                filename=filename,
                rows=len(df),
                numeric_cols=len(numeric_cols),
                total_cols=len(df.columns)
            )

            summary_parts.append("--- MEASUREMENT DATA - Statistical Summary ---")
            summary_parts.append(f"(Detected {len(numeric_cols)}/{len(df.columns)} numeric columns)")
            summary_parts.append("")

            # Statistical summary for each numeric column
            for col in numeric_cols:
                valid_data = df[col].dropna()

                if len(valid_data) == 0:
                    summary_parts.append(f"\n{col}: [No valid data]")
                    continue

                stats = {
                    'count': len(valid_data),
                    'min': valid_data.min(),
                    'max': valid_data.max(),
                    'mean': valid_data.mean(),
                    'median': valid_data.median(),
                    'std': valid_data.std() if len(valid_data) > 1 else 0
                }

                summary_parts.append(f"\n{col}:")
                summary_parts.append(f"  Range: {stats['min']:.6g} to {stats['max']:.6g}")
                summary_parts.append(f"  Mean: {stats['mean']:.6g} (±{stats['std']:.6g} std dev)")
                summary_parts.append(f"  Median: {stats['median']:.6g}")
                summary_parts.append(f"  Valid samples: {stats['count']}/{len(df)}")

            # Add representative samples
            n_samples = min(settings.excel_sample_rows, len(df) // 2)
            if n_samples > 0:
                summary_parts.append("\n--- Representative Samples ---")
                summary_parts.append(f"\nFirst {n_samples} rows:")
                summary_parts.append(df.head(n_samples).to_string(index=False))

                summary_parts.append(f"\nLast {n_samples} rows:")
                summary_parts.append(df.tail(n_samples).to_string(index=False))

        # Strategy 2: Preview for text/categorical data
        else:
            logger.info(
                "using_row_preview",
                filename=filename,
                rows=len(df),
                numeric_cols=len(numeric_cols),
                total_cols=len(df.columns)
            )

            max_rows = min(settings.excel_max_preview_rows, len(df))
            summary_parts.append(f"--- Data Preview (first {max_rows} rows) ---")
            summary_parts.append(df.head(max_rows).to_string(index=False))

            if len(df) > max_rows:
                summary_parts.append(f"\n... ({len(df) - max_rows} more rows not shown)")

        return "\n".join(summary_parts)

    async def _extract_multi_sheet_excel(
        self,
        sheets_dict: dict[str, pd.DataFrame],
        filename: str
    ) -> str:
        """Extract and summarize multi-sheet Excel file."""
        summary_parts = [
            f"Excel file: {filename}",
            f"Total sheets: {len(sheets_dict)}",
            ""
        ]

        for sheet_name, df in sheets_dict.items():
            summary_parts.append(f"\n{'='*70}")
            summary_parts.append(await self._summarize_dataframe(df, filename, sheet_name))
            summary_parts.append('='*70)

        return "\n".join(summary_parts)

    async def _extract_text_from_image_with_vision(
        self,
        image_bytes: bytes,
        page_identifier: str
    ) -> str:
        """
        Extract text from an image using Claude's vision API.

        Args:
            image_bytes: Image data as bytes (PNG format)
            page_identifier: Identifier for logging (e.g., "page_1")

        Returns:
            Extracted text content
        """
        try:
            import google.generativeai as genai
            from app.config import settings

            # Configure Gemini API
            genai.configure(api_key=settings.google_api_key)

            # Use Gemini 2.5 Pro with vision - with retry logic for empty responses
            model = genai.GenerativeModel(settings.gemini_vision_model)

            max_retries = 2
            retry_delay = 2  # seconds

            for attempt in range(max_retries + 1):
                if attempt > 0:
                    logger.warning(
                        "vision_retry_attempt",
                        page=page_identifier,
                        attempt=attempt,
                        max_retries=max_retries
                    )
                    await asyncio.sleep(retry_delay)

                # Prepare image part for Gemini
                image_part = {
                    "mime_type": "image/png",
                    "data": image_bytes
                }

                prompt = """You are a document digitization specialist. Your task is to convert this document page image into a comprehensive, structured JSON format that captures ALL content.

CRITICAL: Extract EVERYTHING visible - miss nothing. This is the only chance to capture this data.

Return a JSON object with this structure:

{
  "document_type": "string (email, technical_drawing, table, form, safety_data_sheet, measurement_report, process_flow_diagram, questionnaire, mixed)",
  "metadata": {
    "has_tables": boolean,
    "has_diagrams": boolean,
    "has_handwriting": boolean,
    "language": "string (de, en, mixed)",
    "page_number": "string or null if visible"
  },
  "content": {
    "body_text": "string (ALL text in document order. CRITICAL: Write headers WITH their paragraphs together, not separately! Format: HEADER\\n\\nparagraph content\\n\\nNEXT HEADER\\n\\nparagraph. Preserve document structure!)",
    "tables": [
      {
        "title": "string or null",
        "headers": ["column1", "column2", ...],
        "rows": [
          ["cell1", "cell2", ...],
          ["cell1", "cell2", ...]
        ]
      }
    ],
    "lists": [
      {
        "type": "bulleted or numbered",
        "items": ["item1", "item2", ...]
      }
    ],
    "key_value_pairs": [
      {"key": "string", "value": "string"}
    ],
    "diagrams_and_images": [
      {
        "type": "flow_diagram, chart, logo, signature, stamp, photo, technical_drawing",
        "description": "detailed description of what is shown",
        "labels_and_text": ["any text visible in/on the diagram"]
      }
    ],
    "signatures_and_stamps": [
      {
        "type": "signature or stamp",
        "text": "any readable text",
        "location": "top_right, bottom_left, etc."
      }
    ]
  },
  "quality_notes": "string (mention any unclear text, cut-off content, poor quality areas)"
}

EXTRACTION RULES:

1. **Tables**:
   - Extract EVERY row and column
   - Preserve exact values, units, symbols (%, ≤, ≥, -, ~)
   - If cells span multiple rows/columns, note this
   - Include table headers AND all data rows

2. **Text**:
   - Extract ALL paragraphs verbatim
   - Preserve line breaks between sections
   - Include page numbers, headers, footers
   - Capture email signatures, contact info

3. **Key-Value Pairs**:
   - Extract form fields: "Field Name: Value"
   - Extract measurement data: "Temperature: 45°C"
   - Extract parameters: "Flow Rate: 5000 m³/h"

4. **Preserve Formatting**:
   - Keep units exactly: m³/h, °C, mg/Nm³, %
   - Keep special characters: ≤, ≥, ±, ~, -, /, ×
   - Keep German umlauts: ä, ö, ü, ß

5. **Diagrams/Images**:
   - Describe what is shown (process flow, equipment layout, etc.)
   - Extract ALL labels, annotations, arrows, text from diagrams
   - Note connections between elements

6. **Don't Miss**:
   - Small print, footnotes, references
   - Handwritten notes or annotations
   - Stamps, signatures, logos with text
   - Section numbers, page numbers
   - CAS numbers, chemical formulas
   - Email headers, sender/recipient info

Return ONLY valid JSON. No markdown, no commentary."""

                # Generate content with Gemini (async)
                response = await model.generate_content_async(
                    [image_part, prompt],
                    generation_config={
                        "temperature": 0.2,
                        "max_output_tokens": 4096,
                    }
                )

                # Check if response is valid (Gemini format)
                if response.text and response.text.strip():
                    # Success! Break retry loop
                    logger.info(
                        "vision_extraction_success",
                        page=page_identifier,
                        length=len(response.text),
                        attempt=attempt
                    )
                    break

                # If we get here, response was empty or invalid
                if attempt < max_retries:
                    logger.warning(
                        "vision_response_empty_retrying",
                        page=page_identifier,
                        attempt=attempt
                    )
                    continue
                else:
                    # Final attempt failed
                    logger.warning(
                        "vision_response_empty_final",
                        page=page_identifier,
                        attempts=max_retries + 1
                    )
                    return f"[No content in vision response for {page_identifier} after {max_retries + 1} attempts]"

            # Safely extract text from response with proper error handling (Gemini format)
            if not response.text or not response.text.strip():
                logger.warning(
                    "vision_response_empty",
                    page=page_identifier
                )
                return f"[No content in vision response for {page_identifier}]"

            extracted_text = response.text

            logger.info(
                "vision_extraction_complete",
                page=page_identifier,
                length=len(extracted_text)
            )

            return extracted_text

        except Exception as e:
            logger.error(
                "vision_extraction_failed",
                page=page_identifier,
                error=str(e)
            )
            return f"[Vision extraction failed for {page_identifier}: {str(e)}]"

    async def _extract_text_from_image_with_openai(
        self,
        image_bytes: bytes,
        page_identifier: str
    ) -> str:
        """
        Extract text from an image using OpenAI GPT-5 vision API.

        Args:
            image_bytes: Image data as bytes (PNG format)
            page_identifier: Identifier for logging (e.g., "page_1")

        Returns:
            Extracted text content
        """
        try:
            import openai
            import asyncio
            from app.config import settings

            # Configure OpenAI API
            client = openai.AsyncOpenAI(api_key=settings.openai_api_key)

            # Use OpenAI GPT-5 with vision - with retry logic for empty responses
            max_retries = 2
            retry_delay = 2  # seconds

            for attempt in range(max_retries + 1):
                if attempt > 0:
                    logger.warning(
                        "vision_retry_attempt_openai",
                        page=page_identifier,
                        attempt=attempt,
                        max_retries=max_retries
                    )
                    await asyncio.sleep(retry_delay)

                # Encode image as base64
                image_base64 = base64.b64encode(image_bytes).decode('utf-8')

                prompt = """You are a document digitization specialist. Your task is to convert this document page image into a comprehensive, structured JSON format that captures ALL content.

CRITICAL: Extract EVERYTHING visible - miss nothing. This is the only chance to capture this data.

Return a JSON object with this structure:

{
  "document_type": "string (email, technical_drawing, table, form, safety_data_sheet, measurement_report, process_flow_diagram, questionnaire, mixed)",
  "metadata": {
    "has_tables": boolean,
    "has_diagrams": boolean,
    "has_handwriting": boolean,
    "language": "string (de, en, mixed)",
    "page_number": "string or null if visible"
  },
  "content": {
    "body_text": "string (ALL text in document order. CRITICAL: Write headers WITH their paragraphs together, not separately! Format: HEADER\\n\\nparagraph content\\n\\nNEXT HEADER\\n\\nparagraph. Preserve document structure!)",
    "tables": [
      {
        "title": "string or null",
        "headers": ["column1", "column2", ...],
        "rows": [
          ["cell1", "cell2", ...],
          ["cell1", "cell2", ...]
        ]
      }
    ],
    "lists": [
      {
        "type": "bulleted or numbered",
        "items": ["item1", "item2", ...]
      }
    ],
    "key_value_pairs": [
      {"key": "string", "value": "string"}
    ],
    "diagrams_and_images": [
      {
        "type": "flow_diagram, chart, logo, signature, stamp, photo, technical_drawing",
        "description": "detailed description of what is shown",
        "labels_and_text": ["any text visible in/on the diagram"]
      }
    ],
    "signatures_and_stamps": [
      {
        "type": "signature or stamp",
        "text": "any readable text",
        "location": "top_right, bottom_left, etc."
      }
    ]
  },
  "quality_notes": "string (mention any unclear text, cut-off content, poor quality areas)"
}

EXTRACTION RULES:

1. **Tables**:
   - Extract EVERY row and column
   - Preserve exact values, units, symbols (%, ≤, ≥, -, ~)
   - If cells span multiple rows/columns, note this
   - Include table headers AND all data rows

2. **Text**:
   - Extract ALL paragraphs verbatim
   - Preserve line breaks between sections
   - Include page numbers, headers, footers
   - Capture email signatures, contact info

3. **Key-Value Pairs**:
   - Extract form fields: "Field Name: Value"
   - Extract measurement data: "Temperature: 45°C"
   - Extract parameters: "Flow Rate: 5000 m³/h"

4. **Preserve Formatting**:
   - Keep units exactly: m³/h, °C, mg/Nm³, %
   - Keep special characters: ≤, ≥, ±, ~, -, /, ×
   - Keep German umlauts: ä, ö, ü, ß

5. **Diagrams/Images**:
   - Describe what is shown (process flow, equipment layout, etc.)
   - Extract ALL labels, annotations, arrows, text from diagrams
   - Note connections between elements

6. **Don't Miss**:
   - Small print, footnotes, references
   - Handwritten notes or annotations
   - Stamps, signatures, logos with text
   - Section numbers, page numbers
   - CAS numbers, chemical formulas
   - Email headers, sender/recipient info

Return ONLY valid JSON. No markdown, no commentary."""

                # Generate content with OpenAI (async)
                # Note: GPT-5 doesn't support temperature parameter
                response = await client.chat.completions.create(
                    model=settings.vision_openai_model,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/png;base64,{image_base64}"
                                    }
                                },
                                {
                                    "type": "text",
                                    "text": prompt
                                }
                            ]
                        }
                    ],
                    max_completion_tokens=16000  # Increased from 4096 - GPT-5 reasoning uses tokens, needs larger budget
                )

                # Debug: Log FULL response structure to understand the issue
                logger.info(
                    "openai_response_debug",
                    page=page_identifier,
                    has_choices=bool(response.choices),
                    num_choices=len(response.choices) if response.choices else 0,
                    finish_reason=response.choices[0].finish_reason if response.choices else None,
                    has_content=bool(response.choices[0].message.content) if response.choices else False,
                    content_length=len(response.choices[0].message.content) if (response.choices and response.choices[0].message.content) else 0,
                    has_refusal=bool(getattr(response.choices[0].message, 'refusal', None)) if response.choices else False,
                    refusal_text=getattr(response.choices[0].message, 'refusal', None) if response.choices else None,
                    usage_dict=response.usage.model_dump() if hasattr(response, 'usage') and response.usage else None
                )

                # Check if response is valid (OpenAI format)
                if response.choices and response.choices[0].message.content:
                    extracted_text = response.choices[0].message.content.strip()
                    if extracted_text:
                        # Success! Break retry loop
                        logger.info(
                            "vision_extraction_success_openai",
                            page=page_identifier,
                            length=len(extracted_text),
                            attempt=attempt
                        )
                        break

                # If we get here, response was empty or invalid
                if attempt < max_retries:
                    logger.warning(
                        "vision_response_empty_retrying_openai",
                        page=page_identifier,
                        attempt=attempt
                    )
                    continue
                else:
                    # Final attempt failed
                    logger.warning(
                        "vision_response_empty_final_openai",
                        page=page_identifier,
                        attempts=max_retries + 1
                    )
                    return f"[No content in vision response for {page_identifier} after {max_retries + 1} attempts]"

            # Safely extract text from response
            if not extracted_text:
                logger.warning(
                    "vision_response_empty_openai",
                    page=page_identifier
                )
                return f"[No content in vision response for {page_identifier}]"

            logger.info(
                "vision_extraction_complete_openai",
                page=page_identifier,
                length=len(extracted_text)
            )

            return extracted_text

        except Exception as e:
            logger.error(
                "vision_extraction_failed_openai",
                page=page_identifier,
                error=str(e)
            )
            return f"[Vision extraction failed for {page_identifier}: {str(e)}]"

    def _blend_vision_results(
        self,
        openai_result: str,
        gemini_result: str,
        page_num: int
    ) -> tuple[str, str]:
        """
        Intelligently blend results from OpenAI and Gemini vision extraction.

        Strategy:
        - Choose the result with more content (character count)
        - Exclude error messages from consideration
        - Log which model was selected

        Args:
            openai_result: Result from OpenAI vision
            gemini_result: Result from Gemini vision
            page_num: Page number for logging

        Returns:
            Tuple of (selected_result, selected_model)
        """
        # Check if either result is an error message
        openai_is_error = openai_result.startswith("[") and openai_result.endswith("]")
        gemini_is_error = gemini_result.startswith("[") and gemini_result.endswith("]")

        # If one is error and other is not, choose the non-error
        if openai_is_error and not gemini_is_error:
            logger.info("blend_chose_gemini", page=page_num, reason="openai_error")
            return gemini_result, "gemini"
        elif gemini_is_error and not openai_is_error:
            logger.info("blend_chose_openai", page=page_num, reason="gemini_error")
            return openai_result, "openai"
        elif openai_is_error and gemini_is_error:
            # Both failed, choose the longer error message (might have more info)
            logger.warning("blend_both_failed", page=page_num)
            if len(openai_result) >= len(gemini_result):
                return openai_result, "openai"
            else:
                return gemini_result, "gemini"

        # Both succeeded, choose based on content length
        openai_len = len(openai_result)
        gemini_len = len(gemini_result)

        if openai_len >= gemini_len:
            logger.info(
                "blend_chose_openai",
                page=page_num,
                openai_len=openai_len,
                gemini_len=gemini_len,
                diff=openai_len - gemini_len
            )
            return openai_result, "openai"
        else:
            logger.info(
                "blend_chose_gemini",
                page=page_num,
                openai_len=openai_len,
                gemini_len=gemini_len,
                diff=gemini_len - openai_len
            )
            return gemini_result, "gemini"

    async def _get_cached_extraction(self, file_path: str) -> Optional[str]:
        """
        Get cached extraction from database if available and file hasn't changed.

        Uses file hash to detect if file has been modified since last extraction.

        Args:
            file_path: Path to the file

        Returns:
            Cached extracted text if available, None otherwise
        """
        try:
            # Calculate file hash
            file_hash = self._calculate_file_hash(file_path)

            async with AsyncSessionLocal() as db:
                # Look up document by file_path
                stmt = select(Document).where(Document.file_path == file_path)
                result = await db.execute(stmt)
                doc = result.scalar_one_or_none()

                if not doc or not doc.extracted_content:
                    return None

                # Check if file hash matches (file hasn't changed)
                cached_hash = doc.extracted_content.get("file_hash")
                if cached_hash != file_hash:
                    logger.info(
                        "extraction_cache_stale",
                        file_path=file_path,
                        reason="file_hash_mismatch"
                    )
                    return None

                # Return cached text
                return doc.extracted_content.get("text")

        except Exception as e:
            logger.warning(
                "extraction_cache_read_failed",
                file_path=file_path,
                error=str(e)
            )
            return None

    async def _cache_extraction(self, file_path: str, extracted_text: str) -> None:
        """
        Cache extraction result in database with file hash.

        Args:
            file_path: Path to the file
            extracted_text: Extracted text content
        """
        try:
            # Calculate file hash
            file_hash = self._calculate_file_hash(file_path)

            async with AsyncSessionLocal() as db:
                # Look up or create document record
                stmt = select(Document).where(Document.file_path == file_path)
                result = await db.execute(stmt)
                doc = result.scalar_one_or_none()

                cache_data = {
                    "text": extracted_text,
                    "file_hash": file_hash,
                    "cached_at": str(Path(file_path).stat().st_mtime)
                }

                if doc:
                    # Update existing document
                    doc.extracted_content = cache_data
                    logger.info(
                        "extraction_cache_updated",
                        file_path=file_path,
                        text_length=len(extracted_text)
                    )
                else:
                    # Note: Cannot create new Document without session_id
                    # Caching only works for documents already in database
                    logger.info(
                        "extraction_cache_skip",
                        file_path=file_path,
                        reason="document_not_in_database"
                    )
                    return

                await db.commit()

        except Exception as e:
            logger.warning(
                "extraction_cache_write_failed",
                file_path=file_path,
                error=str(e)
            )

    def _calculate_file_hash(self, file_path: str) -> str:
        """
        Calculate SHA256 hash of file for cache validation.

        Args:
            file_path: Path to the file

        Returns:
            SHA256 hex digest
        """
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            # Read file in chunks to handle large files
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
